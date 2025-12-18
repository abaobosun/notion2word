"""
HTML 解析和 Word 文档生成模块
将 Notion 页面的 HTML 转换为 Word 文档
"""
import re
import io
import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class NotionToWordConverter:
    """Notion HTML 到 Word 文档转换器"""
    
    def __init__(self):
        """初始化转换器"""
        self.doc = Document()
        self.image_count = 0
    
    def convert(self, html_content: str, output_filename):
        """
        将 HTML 内容转换为 Word 文档
        
        Args:
            html_content: Notion 页面的 HTML 内容
            output_filename: 输出的 Word 文件名或文件对象
            
        Raises:
            Exception: 转换失败
        """
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 查找主内容区域
        content_div = soup.find('div', class_='notion-page-content')
        
        if not content_div:
            raise Exception("无法找到 Notion 内容区域，可能页面未公开或加载失败")
        
        # 提取页面标题
        title = self._extract_title(soup)
        if title:
            self.doc.add_heading(title, level=0)
        
        # 处理所有块
        self._process_blocks(content_div)
        
        # 保存文档
        self.doc.save(output_filename)
    
    def _extract_title(self, soup: BeautifulSoup) -> str:
        """提取页面标题"""
        # 尝试从多个位置提取标题
        title_selectors = [
            'div.notion-page-block h1',
            'h1.notion-header-block',
            'div[data-block-id] h1',
        ]
        
        for selector in title_selectors:
            title_elem = soup.select_one(selector)
            if title_elem:
                return title_elem.get_text(strip=True)
        
        return ""
    
    def _process_blocks(self, content_div: Tag):
        """处理所有内容块"""
        for child in content_div.find_all(recursive=False):
            self._process_block(child)
    
    def _process_block(self, element: Tag):
        """
        处理单个块元素
        
        Args:
            element: BeautifulSoup Tag 对象
        """
        # 跳过空元素
        if not element.name:
            return
        
        # 获取元素的类名
        classes = element.get('class', [])
        class_str = ' '.join(classes) if classes else ''
        
        # 标题块
        if element.name in ['h1', 'h2', 'h3'] or 'notion-header' in class_str:
            self._add_heading(element)
        
        # 段落块
        elif element.name == 'p' or 'notion-text' in class_str or 'notion-callout' in class_str:
            self._add_paragraph(element)
        
        # 列表
        elif element.name in ['ul', 'ol']:
            self._add_list(element)
        
        # 图片
        elif element.name == 'img' or element.find('img'):
            self._add_image(element)
        
        # 代码块
        elif 'notion-code' in class_str or element.name == 'pre':
            self._add_code_block(element)
        
        # 引用块
        elif 'notion-quote' in class_str or element.name == 'blockquote':
            self._add_quote(element)
        
        # 递归处理子元素（如果是容器）
        elif element.find_all(recursive=False):
            for child in element.find_all(recursive=False):
                self._process_block(child)
    
    def _add_heading(self, element: Tag):
        """添加标题"""
        text = element.get_text(strip=True)
        if not text:
            return
        
        # 确定标题级别
        level = 1
        if element.name == 'h1':
            level = 1
        elif element.name == 'h2':
            level = 2
        elif element.name == 'h3':
            level = 3
        else:
            # 根据字体大小判断
            style = element.get('style', '')
            if 'font-size' in style:
                size_match = re.search(r'font-size:\s*(\d+)', style)
                if size_match:
                    size = int(size_match.group(1))
                    if size >= 30:
                        level = 1
                    elif size >= 24:
                        level = 2
                    else:
                        level = 3
        
        self.doc.add_heading(text, level=level)
    
    def _add_paragraph(self, element: Tag):
        """添加段落"""
        text = element.get_text(strip=True)
        if not text:
            return
        
        paragraph = self.doc.add_paragraph()
        self._add_formatted_text(paragraph, element)
    
    def _add_formatted_text(self, paragraph, element: Tag):
        """添加带格式的文本"""
        for content in element.descendants:
            if isinstance(content, str):
                text = content.strip()
                if text:
                    run = paragraph.add_run(text)
                    
                    # 查找父元素的样式
                    parent = content.parent
                    if parent:
                        # 粗体
                        if parent.name in ['strong', 'b']:
                            run.bold = True
                        # 斜体
                        if parent.name in ['em', 'i']:
                            run.italic = True
                        # 下划线
                        if parent.name == 'u':
                            run.underline = True
                        # 代码
                        if parent.name == 'code':
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
    
    def _add_list(self, element: Tag):
        """添加列表"""
        is_ordered = element.name == 'ol'
        style = 'List Number' if is_ordered else 'List Bullet'
        
        for li in element.find_all('li', recursive=False):
            text = li.get_text(strip=True)
            if text:
                self.doc.add_paragraph(text, style=style)
    
    def _add_image(self, element: Tag):
        """添加图片"""
        img = element if element.name == 'img' else element.find('img')
        if not img:
            return
        
        src = img.get('src', '')
        if not src:
            return
        
        try:
            # 下载图片
            if src.startswith('data:image'):
                # Base64 图片
                return  # 暂不处理 Base64
            elif src.startswith('http'):
                response = requests.get(src, timeout=10)
                if response.status_code == 200:
                    image_stream = io.BytesIO(response.content)
                    self.doc.add_picture(image_stream, width=Inches(5))
                    self.image_count += 1
        except Exception as e:
            # 图片加载失败，添加说明文字
            self.doc.add_paragraph(f"[图片加载失败: {str(e)}]")
    
    def _add_code_block(self, element: Tag):
        """添加代码块"""
        code_text = element.get_text(strip=True)
        if not code_text:
            return
        
        paragraph = self.doc.add_paragraph(code_text)
        paragraph.style = 'Quote'
        
        # 设置等宽字体
        for run in paragraph.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
    
    def _add_quote(self, element: Tag):
        """添加引用"""
        text = element.get_text(strip=True)
        if not text:
            return
        
        self.doc.add_paragraph(text, style='Quote')
